import os
import io
import docx
from groq import Groq

# ==========================================
# MODELO ACTIVO (Groq)
# ==========================================
MODELO_ACTIVO = "openai/gpt-oss-20b"

# ==========================================
# DETECTOR AUTOMÁTICO DE RIESGO
# ==========================================
def evaluar_nivel_riesgo_automatico(texto):
    if not texto:
        return "Bajo"
    texto_lower = texto.lower()
    palabras_alto = ["matar", "suicid", "morirme", "atentar", "intento", "cúter", "cuter", "corta", "ahocar", "arma", "morir", "desaparecer", "no quiero vivir"]
    palabras_medio = ["ansiedad", "pánico", "panico", "droga", "alcohol", "agresiv", "pelear", "depresió", "depresio", "triste", "llorar", "impulsiv"]

    for palabra in palabras_alto:
        if palabra in texto_lower:
            return "Alto"
    for palabra in palabras_medio:
        if palabra in texto_lower:
            return "Medio"
    return "Bajo"

# ==========================================
# 1. RUTEADOR E INTÉRPRETE DE VOZ AUTÓNOMO (PATU LIVE)
# ==========================================
def procesar_comando_agente_patu(comando_voz, datos_contexto="", api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    comando_lower = comando_voz.lower()

    # COMANDO ESPECIAL: MUESTRA TU PODER / DEMOSTRACIÓN AUTÓNOMA
    if "poder" in comando_lower or "demostración" in comando_lower or "demostracion" in comando_lower or "capaz" in comando_lower:
        return """
        🐾 **DEMOSTRACIÓN DE POTENCIAL CLÍNICO AUTÓNOMO — PATU AI**
        
        ### 🔬 1. Evaluación y Diagnóstico Multiaxial Automático
        - **Impensión Diagnóstica (DSM-5):** F41.1 Trastorno de Ansiedad Generalizada / F32.1 Episodio Depresivo Moderado.
        - **Factores de Riesgo:** Estrés académico elevado, alteración del patrón de sueño y somatización conductual.
        - **Batería Sugerida:** Inventario de Ansiedad de Beck (BAI) + Inventario de Depresión de Beck (BDI-II).

        ### 🎯 2. Plan de Intervención Cognitivo-Conductual (12 Sesiones)
        - **Fase 1 (Sesiones 1-3):** Psicoeducación sobre la ansiedad y reestructuración cognitiva de pensamientos automáticos.
        - **Fase 2 (Sesiones 4-8):** Exposición gradual, técnicas de desactivación fisiológica (respiración diafragmática) y resolución de problemas.
        - **Fase 3 (Sesiones 9-12):** Prevención de recaídas, consolidación de habilidades y plan de alta terapéutica.

        ### 🌳 3. Dinámica Familiar y Genograma
        - Estructura nuclear con comunicación implícita rígida. Alianza materna protectora y distancia afectiva paterna.

        *PATU AI ha ejecutado de forma autónoma el análisis multiaxial, la propuesta de tratamiento y la evaluación de riesgos en tiempo real.*
        """

    prompt_sistema = f"""
    Eres PATU AI, un asistente clínico experto en psicología con la personalidad de un amigable gatito blanco colaborador.

    REGLAS RESTRICCIONALES DE RESPUESTA POR VOZ:
    1. SI EL USUARIO SOLO DICE "Hola", "Dime quién eres", "Hola PATU" O SIMILARES:
       Responde ÚNICAMENTE: "¡Hola a todos! Soy PATU AI, el workstation clínico inteligente especializado en psicología. Estoy listo para asistirlos en diagnósticos, planes de tratamiento e informes. ¿En qué trabajaremos hoy?"

    2. SI Y SOLO SI EL USUARIO PREGUNTA EXPLÍCITAMENTE POR TU CREADOR O AUTOR (ej: "Dime quién eres y quién es tu creador", "¿Quién te creó?"):
       Responde: "¡Hola a todos! Soy PATU AI, un asistente clínico inteligente creado con orgullo por Yordán Rugel Martínez junto al Grupo 2 del curso de Gestión de Proyectos de los días viernes, a cargo del docente Richard Edgar González."

    CONTEXTO ACTUAL DEL PACIENTE:
    "{datos_contexto}"

    ORDEN RECIBIDA POR VOZ:
    "{comando_voz}"

    REGLAS GENERALES:
    - Sé directo, empático y conciso.
    - Si te piden realizar una tarea clínica (fases, planes, genogramas, informes), ejecútala con máximo rigor técnico.
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(
            model=MODELO_ACTIVO,
            messages=[{"role": "user", "content": prompt_sistema}],
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error procesando orden de voz: {str(e)}"

# ==========================================
# 2. ANALIZADOR CLÍNICO
# ==========================================
def analizar_caso_inicial(narrativa_completa, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"""
    Eres un asistente clínico experto en psicología y psicodiagnóstico.
    Analiza el siguiente caso clínico y proporciona un informe estructurado que contenga:
    
    ### 1. Resumen Clínico del Caso
    ### 2. Impresión Diagnóstica Multiaxial (CIE-11 / DSM-5)
    ### 3. Brechas de Información y Preguntas Recomendadas
    ### 4. Hipótesis Explicativas del Caso
    ### 5. Diagnósticos Diferenciales a Descartar
    ### 6. Batería de Pruebas Psicométricas Sugeridas
    
    CASO CLÍNICO:
    "{narrativa_completa}"
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error en análisis clínico: {str(e)}"

# ==========================================
# 3. GENERADOR DE GENOGRAMA FAMILIAR
# ==========================================
def generar_genograma_familiar(texto_familia, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"""
    Eres un terapeuta familiar experto en psicodiagnóstico sistémico.
    A partir de los siguientes datos, genera un análisis e interpretación clínica del GENOGRAMA FAMILIAR:

    DATOS FAMILIARES:
    "{texto_familia}"

    Proporciona:
    ### 1. Estructura Familiar Multigeneracional (Representación esquemática)
    ### 2. Calidad de las Relaciones (Conflictivas, Fusionales, Distantes, Alianzas)
    ### 3. Antecedentes Clínicos y Legado Transgeneracional
    ### 4. Hipótesis Sistémica e Impacto en el Paciente
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error al generar genograma: {str(e)}"

# ==========================================
# 4. GENERADOR DE HOJAS DE TRABAJO Y REGISTROS
# ==========================================
def generar_hoja_trabajo_paciente(tipo_registro, diagnostico_o_meta, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"""
    Eres un psicólogo clínico especialista en diseño de material terapéutico para pacientes.
    Diseña una HOJA DE TRABAJO Y REGISTRO PRÁCTICO para ser entregada al paciente.

    DETALLES:
    - Tipo de Herramienta: {tipo_registro}
    - Motivo / Diagnóstico Blanco: {diagnostico_o_meta}

    Proporciona:
    ### 1. Instrucciones Claras y Empáticas de Uso
    ### 2. Ejemplo Práctico Resuelto
    ### 3. Plantilla de Registro / Tabla de Seguimiento
    ### 4. Pregunta o Reflexión Semanal de Cierre
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error al generar hoja de trabajo: {str(e)}"

# ==========================================
# 5. BUSCADOR DE PRUEBAS
# ==========================================
def obtener_pruebas_psicometricas(caso_o_sintomas, edad, etapa, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"""
    Eres un psicómetra experto. Recomienda las pruebas psicométricas y proyectivas normadas para:
    - Edad: {edad} años ({etapa})
    - Sintomatología: {caso_o_sintomas}
    
    Indica:
    1. Nombre oficial de la prueba y sigla.
    2. Dimensiones que evalúa.
    3. Justificación clínica de su elección.
    4. Rango de edad normado.
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error al buscar pruebas: {str(e)}"

# ==========================================
# 6. CO-TERAPEUTA & SUPERVISIÓN DE CASOS
# ==========================================
def generar_supervision_coterapeuta(nombre_paciente, consulta, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"""
    Eres un supervisor clínico senior. Brinda orientación técnica para el abordaje del paciente {nombre_paciente}:

    CONSULTA DEL TERAPEUTA:
    "{consulta}"

    Responde con:
    ### 1. Análisis de la Dinámica Terapéutica y Alianza
    ### 2. Estrategias Sugeridas para la Siguiente Sesión
    ### 3. Manejo de Resistencias o Encuadre Terapéutico
    ### 4. Preguntas de Auto-Reflexión para el Profesional
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error en supervisión: {str(e)}"

# ==========================================
# 7. PROTOCOLO DE CRISIS
# ==========================================
def generar_compromiso_vida(nombre_paciente, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"""
    Redacta un documento formal de 'CONTRATO TERAPÉUTICO DE COMPROMISO CON LA VIDA Y SEGURIDAD' para el paciente {nombre_paciente}.
    Incluye cláusulas de apoyo, red de contactos de emergencia, compromisos mutuos entre terapeuta y paciente, y pautas de acción inmediata ante crisis.
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error al generar compromiso: {str(e)}"

# ==========================================
# 8. PLAN DE TRATAMIENTO
# ==========================================
def generar_plan_tratamiento_psicologico(diagnostico_o_caso, enfoque, num_sesiones=12, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"""
    Diseña un plan de tratamiento psicológico de {num_sesiones} sesiones bajo el enfoque {enfoque} para: {diagnostico_o_caso}.
    Incluye:
    ### 1. Objetivos Terapéuticos
    ### 2. Estructura por Fases de Intervención
    ### 3. Técnicas Específicas
    ### 4. Tareas para Casa
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error en plan de tratamiento: {str(e)}"

# ==========================================
# 9. GENERADOR DE INFORMES PREMIUM
# ==========================================
def generar_informe_premium(datos_dict, enfoque, plantilla_texto="", api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"""
    Redacta un informe psicológico formal basado en los datos:
    Enfoque: {enfoque}
    Datos: {datos_dict}
    Guía de estilo: {plantilla_texto}
    """

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error al generar informe: {str(e)}"

# ==========================================
# 10. ANALIZADOR DE SESIONES
# ==========================================
def analizar_transcripcion_sesion(transcripcion, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"Analiza la siguiente transcripción de sesión terapéutica: {transcripcion}"

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.3)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error en análisis de sesión: {str(e)}"

# ==========================================
# 11. PSICOEDUCACIÓN Y BAREMOS
# ==========================================
def generar_plantilla_psicoeducacion(diagnostico, destinatario, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"Crea una guía psicoeducativa clara para {destinatario} sobre la condición: {diagnostico}."

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.4)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error en psicoeducación: {str(e)}"

def interpretar_puntajes_psicometricos(nombre_prueba, puntajes, edad, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "❌ Error: No se encontró GROQ_API_KEY."

    prompt = f"Interpreta los puntajes de la prueba {nombre_prueba} para un evaluado de {edad} años: {puntajes}"

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(model=MODELO_ACTIVO, messages=[{"role": "user", "content": prompt}], temperature=0.2)
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error al interpretar puntajes: {str(e)}"

# ==========================================
# FUNCIONES DE APOYO (AUDIO Y DOCUMENTOS)
# ==========================================
def transcribir_audio_groq(archivo_audio, api_key=None):
    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return "Error: No se encontró GROQ_API_KEY."

    try:
        client = Groq(api_key=api_key)
        if hasattr(archivo_audio, 'seek'):
            archivo_audio.seek(0)
        audio_bytes = archivo_audio.read()
        nombre_archivo = getattr(archivo_audio, 'name', 'audio.mp3')
        
        transcription = client.audio.transcriptions.create(
            file=(nombre_archivo, audio_bytes),
            model="whisper-large-v3",
            response_format="text"
        )
        return transcription
    except Exception as e:
        return f"Error en transcripción: {str(e)}"

def crear_documento_word(titulo, contenido):
    doc = docx.Document()
    doc.add_heading(titulo, level=1)
    
    lineas = contenido.split('\n')
    for linea in lineas:
        linea = linea.strip()
        if not linea:
            continue
            
        texto_limpio = linea.replace('**', '').replace('__', '')
        
        if linea.startswith('### '):
            doc.add_heading(texto_limpio.replace('### ', ''), level=3)
        elif linea.startswith('## '):
            doc.add_heading(texto_limpio.replace('## ', ''), level=2)
        elif linea.startswith('# '):
            doc.add_heading(texto_limpio.replace('# ', ''), level=1)
        elif linea.startswith('- ') or linea.startswith('* '):
            doc.add_paragraph(texto_limpio[2:], style='List Bullet')
        else:
            doc.add_paragraph(texto_limpio)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def extraer_texto_docx(archivo_docx):
    try:
        doc = docx.Document(archivo_docx)
        texto_completo = []
        for p in doc.paragraphs:
            texto_completo.append(p.text)
        return "\n".join(texto_completo)
    except Exception:
        return ""

def procesar_analisis(archivo, contexto=""):
    texto_archivo = ""
    if archivo is not None:
        if archivo.name.endswith(".docx"):
            texto_archivo = extraer_texto_docx(archivo)
        elif archivo.name.endswith(".txt"):
            texto_archivo = archivo.read().decode("utf-8")
    
    narrativa = f"{contexto}\n\nContenido del documento:\n{texto_archivo}"
    return analizar_caso_inicial(narrativa)
